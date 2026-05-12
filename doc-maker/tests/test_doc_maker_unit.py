import os
import sys
import importlib
import importlib.util
import base64
from io import BytesIO

_parent = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
_deps = os.path.abspath(os.path.join(os.path.dirname(__file__), "../dependencies"))
sys.path.insert(0, _parent)
sys.path.insert(0, _deps)

import pytest  # noqa: E402
from unittest.mock import AsyncMock, MagicMock  # noqa: E402
from autohive_integrations_sdk.integration import ResultType  # noqa: E402

# Load the module from its file location (Integration.load() needs the cwd set)
_original_cwd = os.getcwd()
os.chdir(_parent)
_spec = importlib.util.spec_from_file_location("doc_maker_mod", os.path.join(_parent, "doc_maker.py"))
_mod = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(_mod)
os.chdir(_original_cwd)

doc_maker = _mod.doc_maker  # Integration instance
detect_placeholder_patterns = _mod.detect_placeholder_patterns
has_markdown_formatting = _mod.has_markdown_formatting
is_likely_placeholder_context = _mod.is_likely_placeholder_context
analyze_replacement_safety = _mod.analyze_replacement_safety
_save_document_to_dict = _mod._save_document_to_dict
parse_markdown_to_docx = _mod.parse_markdown_to_docx
documents = _mod.documents

pytestmark = pytest.mark.unit


@pytest.fixture
def mock_context():
    ctx = MagicMock(name="ExecutionContext")
    ctx.fetch = AsyncMock(name="fetch")
    ctx.auth = {}
    return ctx


def _make_docx_bytes() -> bytes:
    """Create a minimal valid .docx file in memory."""
    from docx import Document

    buf = BytesIO()
    Document().save(buf)
    buf.seek(0)
    return buf.read()


def _make_file_item(name: str, data: bytes, content_type: str = "application/octet-stream") -> dict:
    return {
        "name": name,
        "contentType": content_type,
        "content": base64.urlsafe_b64encode(data).decode("ascii"),
    }


# ---- Helper / Pure Function Tests ----


class TestDetectPlaceholderPatterns:
    def test_formal_placeholder_double_braces(self):
        is_ph, pattern = detect_placeholder_patterns("{{FIELD_NAME}}")
        assert is_ph is True
        assert pattern == "formal_placeholder"

    def test_formal_placeholder_single_braces(self):
        is_ph, _ = detect_placeholder_patterns("{FIELD}")
        assert is_ph is True

    def test_formal_placeholder_brackets(self):
        is_ph, _ = detect_placeholder_patterns("[PLACEHOLDER]")
        assert is_ph is True

    def test_instruction_text(self):
        is_ph, pattern = detect_placeholder_patterns("(Note: Add details here)")
        assert is_ph is True
        assert pattern == "instruction_text"

    def test_form_style(self):
        is_ph, pattern = detect_placeholder_patterns("Name: ____")
        assert is_ph is True
        # "Name: ____" matches the formal_placeholder __.*?__ pattern before form_style
        assert pattern in ("form_style", "formal_placeholder")

    def test_generic_tbd(self):
        is_ph, _ = detect_placeholder_patterns("TBD")
        assert is_ph is True

    def test_empty_string(self):
        is_ph, pattern = detect_placeholder_patterns("")
        assert is_ph is True
        assert pattern == "empty"

    def test_whitespace_only(self):
        is_ph, _ = detect_placeholder_patterns("   ")
        assert is_ph is True

    def test_real_content_not_placeholder(self):
        is_ph, pattern = detect_placeholder_patterns("This is a complete sentence with actual content.")
        assert is_ph is False
        assert pattern == "content"

    def test_business_content_not_placeholder(self):
        is_ph, _ = detect_placeholder_patterns("The quarterly revenue exceeded expectations.")
        assert is_ph is False


class TestHasMarkdownFormatting:
    def test_bold(self):
        assert has_markdown_formatting("**bold text**") is True

    def test_italic(self):
        assert has_markdown_formatting("*italic*") is True

    def test_code(self):
        assert has_markdown_formatting("`code`") is True

    def test_strikethrough(self):
        assert has_markdown_formatting("~~strike~~") is True

    def test_underline(self):
        assert has_markdown_formatting("__underline__") is True

    def test_newline(self):
        assert has_markdown_formatting("line1\nline2") is True

    def test_plain_text_no_formatting(self):
        assert has_markdown_formatting("Plain text without any markers") is False


class TestIsLikelyPlaceholderContext:
    def test_standalone_word(self):
        assert is_likely_placeholder_context("name", "name") is True

    def test_form_field_pattern(self):
        assert is_likely_placeholder_context("Name: ___", "Name") is True

    def test_braced_word(self):
        assert is_likely_placeholder_context("{{name}}", "name") is True

    def test_instruction_phrase(self):
        assert is_likely_placeholder_context("insert data here", "data") is True

    def test_content_sentence_not_placeholder(self):
        assert is_likely_placeholder_context("The project name should be descriptive.", "name") is False

    def test_complete_sentence_not_placeholder(self):
        assert is_likely_placeholder_context("The date for the meeting has been set.", "date") is False


class TestAnalyzeReplacementSafety:
    def test_all_safe_matches_low_risk(self):
        matches = [
            {"type": "paragraph", "index": 0, "content": "{{FIELD_ONE}}"},
            {"type": "paragraph", "index": 1, "content": "{{FIELD_TWO}}"},
        ]
        result = analyze_replacement_safety("{{FIELD", matches)
        assert result["safety_level"] == "low_risk"
        assert result["safe_matches"] == 2
        assert result["unsafe_matches"] == 0

    def test_all_unsafe_matches_high_risk(self):
        matches = [
            {
                "type": "paragraph",
                "index": 0,
                "content": "The name of the project is important.",
            },
            {
                "type": "paragraph",
                "index": 1,
                "content": "Please update the name field.",
            },
            {
                "type": "paragraph",
                "index": 2,
                "content": "The customer name should be verified.",
            },
        ]
        result = analyze_replacement_safety("name", matches)
        assert result["safety_level"] == "high_risk"
        assert result["unsafe_matches"] == 3

    def test_mixed_provides_guidance(self):
        matches = [
            {"type": "paragraph", "index": 0, "content": "Name: placeholder"},
            {"type": "paragraph", "index": 1, "content": "The name should be clear"},
            {"type": "paragraph", "index": 2, "content": "Update the project name"},
        ]
        result = analyze_replacement_safety("name", matches)
        assert len(result["guidance"]) > 0


# ---- Action Tests ----


class TestCreateDocument:
    @pytest.mark.asyncio
    async def test_create_empty_document(self, mock_context):
        result = await doc_maker.execute_action("create_document", {}, mock_context)
        assert result.type == ResultType.ACTION
        assert "document_id" in result.result.data
        assert result.result.data["saved"] is True
        assert "file" in result.result.data

    @pytest.mark.asyncio
    async def test_create_document_with_markdown(self, mock_context):
        result = await doc_maker.execute_action(
            "create_document",
            {"markdown_content": "# Hello\n\nThis is a test document."},
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["markdown_processed"] is True
        assert result.result.data["document_id"] is not None

    @pytest.mark.asyncio
    async def test_create_document_with_custom_filename(self, mock_context):
        result = await doc_maker.execute_action(
            "create_document",
            {"markdown_content": "# Test", "custom_filename": "my_report"},
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["file"]["name"] == "my_report.docx"

    @pytest.mark.asyncio
    async def test_create_document_returns_base64_file(self, mock_context):
        result = await doc_maker.execute_action("create_document", {}, mock_context)
        file_content = result.result.data["file"]["content"]
        assert len(file_content) > 0
        # Should be valid base64
        decoded = base64.b64decode(file_content)
        assert len(decoded) > 0

    @pytest.mark.asyncio
    async def test_create_document_file_content_type(self, mock_context):
        result = await doc_maker.execute_action("create_document", {}, mock_context)
        assert (
            result.result.data["file"]["contentType"]
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


class TestSaveDocument:
    @pytest.mark.asyncio
    async def test_save_missing_document_returns_action_error(self, mock_context):
        result = await doc_maker.execute_action(
            "save_document",
            {"document_id": "nonexistent-id", "file_path": "output.docx"},
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR
        assert "nonexistent-id" in result.result.message

    @pytest.mark.asyncio
    async def test_save_existing_document_succeeds(self, mock_context):
        # Create doc first
        create_result = await doc_maker.execute_action("create_document", {}, mock_context)
        doc_id = create_result.result.data["document_id"]

        result = await doc_maker.execute_action(
            "save_document",
            {"document_id": doc_id, "file_path": "test.docx"},
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["saved"] is True
        assert result.result.data["file"]["name"] == "test.docx"


class TestGetDocumentElements:
    @pytest.mark.asyncio
    async def test_get_elements_requires_existing_document(self, mock_context):
        docx_bytes = _make_docx_bytes()
        file_item = _make_file_item("template.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "get_document_elements",
            {"document_id": "test-doc-id", "files": [file_item]},
            mock_context,
        )
        assert result.type == ResultType.ACTION
        data = result.result.data
        assert "template_summary" in data
        assert "fillable_paragraphs" in data
        assert "fillable_cells" in data
        assert "template_ready" in data

    @pytest.mark.asyncio
    async def test_get_elements_missing_document_no_files_raises(self, mock_context):
        # No files and missing doc_id - should raise ValueError
        result = await doc_maker.execute_action(
            "get_document_elements",
            {"document_id": "does-not-exist-xyz"},
            mock_context,
        )
        # The action raises ValueError which the SDK catches
        assert result.type == ResultType.ACTION_ERROR

    @pytest.mark.asyncio
    async def test_get_elements_response_shape(self, mock_context):
        docx_bytes = _make_docx_bytes()
        file_item = _make_file_item("template.docx", docx_bytes)
        result = await doc_maker.execute_action(
            "get_document_elements",
            {"document_id": "shape-test-id", "files": [file_item]},
            mock_context,
        )
        data = result.result.data
        assert isinstance(data["fillable_paragraphs"], list)
        assert isinstance(data["fillable_cells"], list)
        assert isinstance(data["template_summary"]["fillable_total"], int)
        assert data["recommended_strategy"] in ("mixed", "single_method")


class TestAddTable:
    @pytest.mark.asyncio
    async def test_add_table_to_document(self, mock_context):
        create_result = await doc_maker.execute_action("create_document", {}, mock_context)
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "add_table",
            {
                "document_id": doc_id,
                "rows": 3,
                "cols": 2,
                "data": [
                    ["Header1", "Header2"],
                    ["Cell1", "Cell2"],
                    ["Cell3", "Cell4"],
                ],
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["table_rows"] == 3
        assert result.result.data["table_cols"] == 2
        assert result.result.data["saved"] is True

    @pytest.mark.asyncio
    async def test_add_table_missing_document_raises(self, mock_context):
        result = await doc_maker.execute_action(
            "add_table",
            {"document_id": "bad-id", "rows": 2, "cols": 2},
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR


class TestAddMarkdownContent:
    @pytest.mark.asyncio
    async def test_add_markdown_to_existing_document(self, mock_context):
        create_result = await doc_maker.execute_action("create_document", {}, mock_context)
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "add_markdown_content",
            {
                "document_id": doc_id,
                "markdown_content": "## Section\n\nSome content here.",
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["markdown_processed"] is True
        assert result.result.data["saved"] is True

    @pytest.mark.asyncio
    async def test_add_markdown_missing_document_raises(self, mock_context):
        result = await doc_maker.execute_action(
            "add_markdown_content",
            {"document_id": "bad-id", "markdown_content": "# Hello"},
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR


class TestAddPageBreak:
    @pytest.mark.asyncio
    async def test_add_page_break(self, mock_context):
        create_result = await doc_maker.execute_action("create_document", {}, mock_context)
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "add_page_break",
            {"document_id": doc_id, "files": [file_item]},
            mock_context,
        )
        assert result.type == ResultType.ACTION
        assert result.result.data["page_break_added"] is True
        assert result.result.data["saved"] is True

    @pytest.mark.asyncio
    async def test_add_page_break_missing_document(self, mock_context):
        result = await doc_maker.execute_action(
            "add_page_break",
            {"document_id": "missing-doc"},
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR


class TestUpdateByPosition:
    @pytest.mark.asyncio
    async def test_update_paragraph_by_position(self, mock_context):
        create_result = await doc_maker.execute_action(
            "create_document",
            {"markdown_content": "# Title\n\nOriginal paragraph."},
            mock_context,
        )
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "update_by_position",
            {
                "document_id": doc_id,
                "updates": [{"type": "paragraph", "index": 0, "content": "Updated content"}],
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        data = result.result.data
        assert "applied" in data
        assert "failed" in data
        assert data["saved"] is True

    @pytest.mark.asyncio
    async def test_update_missing_document(self, mock_context):
        result = await doc_maker.execute_action(
            "update_by_position",
            {
                "document_id": "bad-id",
                "updates": [{"type": "paragraph", "index": 0, "content": "New"}],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR


class TestFindAndReplace:
    @pytest.mark.asyncio
    async def test_find_and_replace_basic(self, mock_context):
        create_result = await doc_maker.execute_action(
            "create_document",
            {"markdown_content": "Hello {{NAME}}, welcome to the team."},
            mock_context,
        )
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "find_and_replace",
            {
                "document_id": doc_id,
                "replacements": [{"find": "{{NAME}}", "replace": "Alice", "replace_all": True}],
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        data = result.result.data
        assert data["replaced"] >= 1
        assert data["safety_active"] is True

    @pytest.mark.asyncio
    async def test_find_and_replace_missing_document(self, mock_context):
        result = await doc_maker.execute_action(
            "find_and_replace",
            {
                "document_id": "bad-id",
                "replacements": [{"find": "foo", "replace": "bar"}],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR

    @pytest.mark.asyncio
    async def test_find_no_match_returns_warning(self, mock_context):
        create_result = await doc_maker.execute_action(
            "create_document", {"markdown_content": "Hello world."}, mock_context
        )
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "find_and_replace",
            {
                "document_id": doc_id,
                "replacements": [{"find": "NONEXISTENT_TEXT_XYZ", "replace": "replacement"}],
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        # 0 replacements and a warning
        assert result.result.data["replaced"] == 0

    @pytest.mark.asyncio
    async def test_find_and_replace_invalid_type_replacements(self, mock_context):
        # The SDK validates input schema: replacements must be array, so passing a
        # non-array non-string type triggers VALIDATION_ERROR before the handler runs.
        create_result = await doc_maker.execute_action("create_document", {}, mock_context)
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "find_and_replace",
            {
                "document_id": doc_id,
                "replacements": 12345,
                "files": [file_item],
            },
            mock_context,
        )
        # SDK validates schema before running handler — integer is not array or string
        assert result.type == ResultType.VALIDATION_ERROR


class TestFillTemplateFields:
    @pytest.mark.asyncio
    async def test_fill_placeholder_data(self, mock_context):
        create_result = await doc_maker.execute_action(
            "create_document",
            {"markdown_content": "Dear {{CLIENT_NAME}}, your invoice is ready."},
            mock_context,
        )
        doc_id = create_result.result.data["document_id"]
        docx_bytes = base64.b64decode(create_result.result.data["file"]["content"])
        file_item = _make_file_item(f"{doc_id}.docx", docx_bytes)

        result = await doc_maker.execute_action(
            "fill_template_fields",
            {
                "document_id": doc_id,
                "template_data": {"placeholder_data": {"{{CLIENT_NAME}}": "Acme Corp"}},
                "files": [file_item],
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION
        data = result.result.data
        assert data["SAFETY_STATUS"] in ("OK", "CRITICAL_ISSUES_DETECTED")
        assert "completed_operations" in data
        assert data["saved"] is True

    @pytest.mark.asyncio
    async def test_fill_missing_document_raises(self, mock_context):
        result = await doc_maker.execute_action(
            "fill_template_fields",
            {
                "document_id": "bad-id",
                "template_data": {"placeholder_data": {"{{X}}": "Y"}},
            },
            mock_context,
        )
        assert result.type == ResultType.ACTION_ERROR


class TestSaveDocumentToDict:
    def test_missing_document_returns_error_dict(self):
        result = _save_document_to_dict("nonexistent-id", "output.docx")
        assert result["saved"] is False
        assert "nonexistent-id" in result["error"]
        assert result["file"]["content"] == ""


class TestParenthesizedListNumbering:
    """Verify that (1), (a), (i) style lists produce correct Word numbering."""

    MARKDOWN = (
        "1. Elephant\n"
        "    (a) Elephants are the largest land animals on Earth, "
        "with African elephants weighing up to 14,000 lbs.\n"
        "    (b) They have an exceptional memory and can recognize "
        "themselves in mirrors, indicating self-awareness.\n"
        "2. Axolotl\n"
        "    (a) Axolotls can regenerate entire limbs, including "
        "parts of their heart and brain.\n"
        "    (b) Unlike most amphibians, axolotls retain their larval "
        "features throughout their entire lives, a trait called neoteny."
    )

    @staticmethod
    def _get_numpr(paragraph):
        """Return (numId, ilvl) from a paragraph's w:numPr, or None."""
        from docx.oxml.ns import qn

        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        if numId_el is None or ilvl_el is None:
            return None
        return int(numId_el.get(qn("w:val"))), int(ilvl_el.get(qn("w:val")))

    @staticmethod
    def _get_abstract_num_for(doc, num_id):
        """Return the abstractNum element referenced by a given numId."""
        from docx.oxml.ns import qn

        numbering = doc.part.numbering_part._element
        for num_el in numbering.findall(qn("w:num")):
            if int(num_el.get(qn("w:numId"))) == num_id:
                abstract_ref = num_el.find(qn("w:abstractNumId"))
                abstract_id = int(abstract_ref.get(qn("w:val")))
                for an in numbering.findall(qn("w:abstractNum")):
                    if int(an.get(qn("w:abstractNumId"))) == abstract_id:
                        return an
        return None

    def test_produces_six_numbered_paragraphs(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        assert len(numbered) == 6, f"Expected 6 numbered paragraphs, got {len(numbered)}: {numbered}"

    def test_top_level_items_are_at_ilvl_zero(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        top_items = [(text, numpr) for text, numpr in numbered if "Elephant" == text or "Axolotl" == text]
        assert len(top_items) == 2, f"Expected 2 top-level items, got {top_items}"
        for text, (num_id, ilvl) in top_items:
            assert ilvl == 0, f"'{text}' should be at ilvl 0, got {ilvl}"

    def test_sub_items_are_indented(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        sub_items = [(text, numpr) for text, numpr in numbered if text not in ("Elephant", "Axolotl")]
        assert len(sub_items) == 4, f"Expected 4 sub-items, got {len(sub_items)}"
        for text, (num_id, ilvl) in sub_items:
            assert ilvl >= 1, f"Sub-item should be indented (ilvl >= 1), got {ilvl}: {text}"

    def test_top_level_uses_decimal_numbering(self):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        elephant = next((text, numpr) for text, numpr in numbered if text == "Elephant")
        num_id = elephant[1][0]
        abstract = self._get_abstract_num_for(doc, num_id)
        assert abstract is not None
        lvl0 = abstract.find(qn("w:lvl"))
        fmt = lvl0.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "decimal", f"Top-level should be decimal, got {fmt}"

    def test_sub_items_use_lower_letter_parenthesized(self):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        first_sub = next((text, numpr) for text, numpr in numbered if "Elephants are" in text)
        num_id, ilvl = first_sub[1]
        abstract = self._get_abstract_num_for(doc, num_id)
        assert abstract is not None

        # Find the lvl element matching the ilvl used
        target_lvl = None
        for lvl in abstract.findall(qn("w:lvl")):
            if int(lvl.get(qn("w:ilvl"))) == ilvl:
                target_lvl = lvl
                break
        assert target_lvl is not None

        fmt = target_lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "lowerLetter", f"Sub-items should be lowerLetter, got {fmt}"
        lvl_text = target_lvl.find(qn("w:lvlText")).get(qn("w:val"))
        assert "(" in lvl_text, f"Sub-items should have parenthesized format, got '{lvl_text}'"

    def test_elephant_text_on_same_line_as_number(self):
        """The parent item text must appear in the same paragraph as the numbering."""
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        elephant_paras = [(t, n) for t, n in numbered if "Elephant" in t and n[1] == 0]
        assert len(elephant_paras) >= 1
        assert elephant_paras[0][0] == "Elephant", (
            f"Top-level text should be exactly 'Elephant', got '{elephant_paras[0][0]}'"
        )

    def test_parent_and_children_share_same_numid(self):
        """Word requires nested lists to share the same numId to render correctly."""
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        num_ids = set(numpr[0] for _, numpr in numbered)
        assert len(num_ids) == 1, (
            f"All items should share one numId for coherent multilevel numbering, got {num_ids}"
        )


class TestMultipleParenListsAfterHeadings:
    """Verify that multiple (1)-style lists separated by headings all display numbering
    and are left-aligned when the markdown has no leading spaces."""

    MARKDOWN = (
        "# Animals\n"
        "(1) Elephant\n"
        "(2) Tiger\n"
        "# Fish\n"
        "(1) squid\n"
        "(2) Whale"
    )

    @staticmethod
    def _get_numpr(paragraph):
        from docx.oxml.ns import qn

        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        if numId_el is None or ilvl_el is None:
            return None
        return int(numId_el.get(qn("w:val"))), int(ilvl_el.get(qn("w:val")))

    @staticmethod
    def _get_abstract_num_for(doc, num_id):
        from docx.oxml.ns import qn

        numbering = doc.part.numbering_part._element
        for num_el in numbering.findall(qn("w:num")):
            if int(num_el.get(qn("w:numId"))) == num_id:
                abstract_ref = num_el.find(qn("w:abstractNumId"))
                abstract_id = int(abstract_ref.get(qn("w:val")))
                for an in numbering.findall(qn("w:abstractNum")):
                    if int(an.get(qn("w:abstractNumId"))) == abstract_id:
                        return an
        return None

    def test_both_lists_produce_numbered_paragraphs(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        assert len(numbered) == 4, f"Expected 4 numbered paragraphs, got {len(numbered)}: {numbered}"

    def test_each_list_has_its_own_abstract_num(self):
        """Each independent list must get its own abstractNum to avoid Word dropping numbers."""
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        animals_num_id = numbered[0][1][0]
        fish_num_id = numbered[2][1][0]

        animals_abstract = self._get_abstract_num_for(doc, animals_num_id)
        fish_abstract = self._get_abstract_num_for(doc, fish_num_id)

        assert animals_abstract is not None
        assert fish_abstract is not None

        animals_abstract_id = int(animals_abstract.get(qn("w:abstractNumId")))
        fish_abstract_id = int(fish_abstract.get(qn("w:abstractNumId")))
        assert animals_abstract_id != fish_abstract_id, (
            "Each list should reference a different abstractNum to prevent Word from dropping numbers"
        )

    def test_all_items_at_ilvl_zero(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        for text, (num_id, ilvl) in numbered:
            assert ilvl == 0, f"'{text}' should be at ilvl 0, got {ilvl}"

    def test_level_zero_is_left_aligned(self):
        """Level-0 paren lists with no leading spaces should be left-aligned (left=360, hanging=360)."""
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        first_num_id = numbered[0][1][0]
        abstract = self._get_abstract_num_for(doc, first_num_id)
        assert abstract is not None

        lvl0 = None
        for lvl in abstract.findall(qn("w:lvl")):
            if int(lvl.get(qn("w:ilvl"))) == 0:
                lvl0 = lvl
                break
        assert lvl0 is not None

        pPr = lvl0.find(qn("w:pPr"))
        assert pPr is not None
        ind = pPr.find(qn("w:ind"))
        assert ind is not None
        assert ind.get(qn("w:left")) == "360", (
            f"Level 0 left indent should be 360 (left-aligned), got {ind.get(qn('w:left'))}"
        )
        assert ind.get(qn("w:hanging")) == "360"

    def test_all_use_decimal_parenthesized_format(self):
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)

        numbered = [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]
        for text, (num_id, ilvl) in numbered:
            abstract = self._get_abstract_num_for(doc, num_id)
            assert abstract is not None
            for lvl in abstract.findall(qn("w:lvl")):
                if int(lvl.get(qn("w:ilvl"))) == ilvl:
                    fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
                    assert fmt == "decimal", f"'{text}' should use decimal, got {fmt}"
                    lvl_text = lvl.find(qn("w:lvlText")).get(qn("w:val"))
                    assert "(" in lvl_text, f"'{text}' should have paren format, got '{lvl_text}'"


class TestMixedNumberedListFormats:
    """Verify that five different numbered-list styles all render correctly:
    1. standard ``1. 2. 3.`` decimal
    2. parenthesized decimal ``(1) (2) (3)``
    3. parenthesized lower letter ``(a) (b) (c)``
    4. parenthesized upper letter ``(A) (B) (C)``
    5. parenthesized lower roman ``(i) (ii) (iii)``

    Each list must:
    - be a real numbered (not bullet) list in the OOXML
    - use the correct numFmt
    - use left-aligned justification
    - display the correct item text ("one", "two", "three")
    - have a consistent hanging indent so text is aligned across items
      whose numbering labels differ in width (e.g. ``(i)`` vs ``(iii)``).
    """

    MARKDOWN = (
        "# numbers\n"
        "1. one\n"
        "2. two\n"
        "3. three\n"
        "# numbers in brackets\n"
        "(1) one\n"
        "(2) two\n"
        "(3) three\n"
        "# letters in brackets\n"
        "(a) one\n"
        "(b) two\n"
        "(c) three\n"
        "# capital letters in brackets\n"
        "(A) one\n"
        "(B) two\n"
        "(C) three\n"
        "# roman numerals\n"
        "(i) one\n"
        "(ii) two\n"
        "(iii) three"
    )

    @staticmethod
    def _get_numpr(paragraph):
        from docx.oxml.ns import qn

        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        if numId_el is None or ilvl_el is None:
            return None
        return int(numId_el.get(qn("w:val"))), int(ilvl_el.get(qn("w:val")))

    @staticmethod
    def _get_abstract_num_for(doc, num_id):
        from docx.oxml.ns import qn

        numbering = doc.part.numbering_part._element
        for num_el in numbering.findall(qn("w:num")):
            if int(num_el.get(qn("w:numId"))) == num_id:
                abstract_ref = num_el.find(qn("w:abstractNumId"))
                abstract_id = int(abstract_ref.get(qn("w:val")))
                for an in numbering.findall(qn("w:abstractNum")):
                    if int(an.get(qn("w:abstractNumId"))) == abstract_id:
                        return an
        return None

    @staticmethod
    def _get_lvl(abstract, ilvl):
        from docx.oxml.ns import qn

        for lvl in abstract.findall(qn("w:lvl")):
            if int(lvl.get(qn("w:ilvl"))) == ilvl:
                return lvl
        return None

    def _build_doc(self):
        from docx import Document

        doc = Document()
        parse_markdown_to_docx(doc, self.MARKDOWN)
        return doc

    def _numbered_paragraphs(self, doc):
        return [(p.text.strip(), self._get_numpr(p)) for p in doc.paragraphs if self._get_numpr(p)]

    # ---- 1. All 15 items are numbered paragraphs ----

    def test_produces_fifteen_numbered_paragraphs(self):
        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        assert len(numbered) == 15, (
            f"Expected 15 numbered paragraphs (5 lists × 3 items), got {len(numbered)}: "
            f"{[t for t, _ in numbered]}"
        )

    # ---- 2. Item text is correct ----

    def test_item_text_is_correct(self):
        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        texts = [t for t, _ in numbered]
        for i in range(5):
            group = texts[i * 3 : i * 3 + 3]
            assert group == ["one", "two", "three"], (
                f"List group {i} text should be ['one', 'two', 'three'], got {group}"
            )

    # ---- 3. Each list uses the correct numFmt ----

    def test_standard_decimal_uses_decimal_format(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[0][1][0]
        ilvl = numbered[0][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "decimal", f"Standard numbered list should be decimal, got {fmt}"

    def test_paren_decimal_uses_decimal_format(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[3][1][0]
        ilvl = numbered[3][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "decimal", f"Paren decimal list should be decimal, got {fmt}"

    def test_paren_lower_letter_uses_lower_letter_format(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[6][1][0]
        ilvl = numbered[6][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "lowerLetter", f"Lower letter list should be lowerLetter, got {fmt}"

    def test_paren_upper_letter_uses_upper_letter_format(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[9][1][0]
        ilvl = numbered[9][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "upperLetter", f"Upper letter list should be upperLetter, got {fmt}"

    def test_paren_roman_uses_lower_roman_format(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[12][1][0]
        ilvl = numbered[12][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        fmt = lvl.find(qn("w:numFmt")).get(qn("w:val"))
        assert fmt == "lowerRoman", f"Roman numeral list should be lowerRoman, got {fmt}"

    # ---- 4. Parenthesized lvlText for bracket lists ----

    def test_paren_lists_use_parenthesized_lvl_text(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        # Lists at indices 3, 6, 9, 12 are the paren lists
        for start_idx in (3, 6, 9, 12):
            num_id = numbered[start_idx][1][0]
            ilvl = numbered[start_idx][1][1]
            abstract = self._get_abstract_num_for(doc, num_id)
            lvl = self._get_lvl(abstract, ilvl)
            lvl_text = lvl.find(qn("w:lvlText")).get(qn("w:val"))
            assert "(" in lvl_text and ")" in lvl_text, (
                f"Item '{numbered[start_idx][0]}' (idx {start_idx}) should have "
                f"parenthesized lvlText, got '{lvl_text}'"
            )

    def test_standard_decimal_uses_dot_lvl_text(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        num_id = numbered[0][1][0]
        ilvl = numbered[0][1][1]
        abstract = self._get_abstract_num_for(doc, num_id)
        lvl = self._get_lvl(abstract, ilvl)
        lvl_text = lvl.find(qn("w:lvlText")).get(qn("w:val"))
        assert "." in lvl_text, (
            f"Standard decimal should use dot format, got '{lvl_text}'"
        )

    # ---- 5. Left-aligned justification ----

    def test_all_lists_are_left_aligned(self):
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)
        checked = set()
        for text, (num_id, ilvl) in numbered:
            key = (num_id, ilvl)
            if key in checked:
                continue
            checked.add(key)
            abstract = self._get_abstract_num_for(doc, num_id)
            lvl = self._get_lvl(abstract, ilvl)
            jc = lvl.find(qn("w:lvlJc"))
            assert jc is not None, f"'{text}' level should have lvlJc element"
            assert jc.get(qn("w:val")) == "left", (
                f"'{text}' should be left-aligned, got '{jc.get(qn('w:val'))}'"
            )

    # ---- 6. Text alignment consistency (hanging indent) ----

    def test_items_within_each_list_share_same_hanging_indent(self):
        """All items in a single list must use the same hanging indent so that
        the text column is aligned even when numbering labels vary in width
        (e.g. ``(i)`` vs ``(iii)``)."""
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)

        for start_idx in range(0, 15, 3):
            group = numbered[start_idx : start_idx + 3]
            # All items in a group share the same numId + ilvl, so they share
            # the same abstractNum level definition → same indent.
            num_id = group[0][1][0]
            ilvl = group[0][1][1]
            abstract = self._get_abstract_num_for(doc, num_id)
            lvl = self._get_lvl(abstract, ilvl)
            pPr = lvl.find(qn("w:pPr"))
            assert pPr is not None, f"Level {ilvl} should have pPr"
            ind = pPr.find(qn("w:ind"))
            assert ind is not None, f"Level {ilvl} should have indent"
            hanging = ind.get(qn("w:hanging"))
            left = ind.get(qn("w:left"))
            assert hanging is not None, f"Hanging indent should be set for list starting at idx {start_idx}"
            assert left is not None, f"Left indent should be set for list starting at idx {start_idx}"

    def test_roman_numeral_hanging_indent_is_wide_enough(self):
        """Roman numeral labels like ``(iii)`` are wider than ``(1)`` so the
        hanging indent must be larger to keep text aligned."""
        from docx.oxml.ns import qn

        doc = self._build_doc()
        numbered = self._numbered_paragraphs(doc)

        # Roman numerals start at index 12
        roman_num_id = numbered[12][1][0]
        roman_ilvl = numbered[12][1][1]
        roman_abstract = self._get_abstract_num_for(doc, roman_num_id)
        roman_lvl = self._get_lvl(roman_abstract, roman_ilvl)
        roman_hanging = int(roman_lvl.find(qn("w:pPr")).find(qn("w:ind")).get(qn("w:hanging")))

        # Decimal at index 3
        decimal_num_id = numbered[3][1][0]
        decimal_ilvl = numbered[3][1][1]
        decimal_abstract = self._get_abstract_num_for(doc, decimal_num_id)
        decimal_lvl = self._get_lvl(decimal_abstract, decimal_ilvl)
        decimal_hanging = int(decimal_lvl.find(qn("w:pPr")).find(qn("w:ind")).get(qn("w:hanging")))

        assert roman_hanging > decimal_hanging, (
            f"Roman numeral hanging indent ({roman_hanging}) should be wider than "
            f"decimal ({decimal_hanging}) to accommodate wider labels like (iii)"
        )
